"""
DOCX Transcript Parser
Extracts transcript data from Word (.docx) files using python-docx.
"""

import re
from typing import Optional, List, Dict
from pathlib import Path

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from models.transcript import Transcript, CourseRecord


class DOCXParser:
    """
    Parses DOCX transcript files into Transcript objects.

    Handles:
    1. Table-based DOCX (extracts from Word tables)
    2. Text-based DOCX (fallback regex extraction from paragraphs)
    """

    COURSE_CODE_RE = re.compile(r"\b([A-Z]{2,4}\d{3,4}[A-Z]?)\b")
    GRADE_RE = re.compile(r"\b(A|A-|B\+|B|B-|C\+|C|C-|D\+|D|F|W|I)\b")
    SEMESTER_RE = re.compile(
        r"\b(Spring|Summer|Fall|Winter)\s+(\d{4})\b", re.IGNORECASE
    )

    HEADER_PATTERNS = {
        "course_code": re.compile(r"course[\s_]*(?:code|id|no)?", re.IGNORECASE),
        "credits": re.compile(r"credit(?:s)?|cr\.?|hours?", re.IGNORECASE),
        "grade": re.compile(r"grade|letter[\s_]*grade|mark", re.IGNORECASE),
        "semester": re.compile(r"semester|term|session|period", re.IGNORECASE),
    }

    @classmethod
    def parse_docx(
        cls,
        file_path: str,
        student_id: Optional[str] = None,
        student_name: Optional[str] = None,
        program: Optional[str] = None,
    ) -> Transcript:
        """
        Parse a DOCX transcript file.

        Args:
            file_path: Path to DOCX file
            student_id: Optional student ID
            student_name: Optional student name
            program: Optional program name

        Returns:
            Transcript object

        Raises:
            ImportError: If python-docx is not installed
            FileNotFoundError: If file doesn't exist
            ValueError: If no valid records found
        """
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX parsing.\n"
                "Install it with: pip install python-docx"
            )

        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        transcript = Transcript(
            student_id=student_id,
            student_name=student_name,
            program=program,
        )

        doc = DocxDocument(file_path)

        # Try table extraction first
        records = cls._extract_from_tables(doc)

        # Fall back to paragraph text
        if not records:
            records = cls._extract_from_paragraphs(doc)

        if not records:
            raise ValueError(
                f"Could not extract transcript data from {file_path}.\n"
                f"Ensure the DOCX contains a table with columns:\n"
                f"  Course Code, Credits, Grade, Semester"
            )

        for record in records:
            transcript.add_record(record)

        return transcript

    @classmethod
    def _extract_from_tables(cls, doc) -> List[CourseRecord]:
        """Extract records from DOCX tables."""
        records = []

        for table in doc.tables:
            rows = table.rows
            if len(rows) < 2:
                continue

            # Try to identify columns from header row
            header_cells = [cell.text.strip() for cell in rows[0].cells]
            col_map = cls._identify_columns(header_cells)

            if col_map:
                for row in rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    record = cls._parse_table_row(cells, col_map)
                    if record:
                        records.append(record)
            else:
                # Try positional parsing
                for row in rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    record = cls._parse_row_positional(cells)
                    if record:
                        records.append(record)

        return records

    @classmethod
    def _extract_from_paragraphs(cls, doc) -> List[CourseRecord]:
        """Fallback: Extract records from paragraph text."""
        records = []
        current_semester = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check for semester header
            sem_match = cls.SEMESTER_RE.search(text)
            if sem_match:
                current_semester = f"{sem_match.group(1).title()} {sem_match.group(2)}"

            # Check for course record
            course_match = cls.COURSE_CODE_RE.search(text)
            grade_match = cls.GRADE_RE.search(text)

            if course_match and grade_match:
                course_code = course_match.group(1)
                grade = grade_match.group(1)

                # Try to extract credits
                credits = 3.0
                numbers = re.findall(r"\b(\d+(?:\.\d+)?)\b", text)
                for num_str in numbers:
                    val = float(num_str)
                    if 0 <= val <= 6:
                        credits = val
                        break

                semester = current_semester or "Unknown"
                line_sem = cls.SEMESTER_RE.search(text)
                if line_sem:
                    semester = f"{line_sem.group(1).title()} {line_sem.group(2)}"

                try:
                    record = CourseRecord(
                        course_code=course_code,
                        credits=credits,
                        grade=grade,
                        semester=semester,
                    )
                    records.append(record)
                except ValueError:
                    continue

        return records

    @classmethod
    def _identify_columns(cls, header_row: list) -> Optional[Dict[str, int]]:
        """Try to identify column indices from a header row."""
        col_map = {}
        for idx, cell in enumerate(header_row):
            if not cell:
                continue
            for field, pattern in cls.HEADER_PATTERNS.items():
                if pattern.search(cell) and field not in col_map:
                    col_map[field] = idx
                    break

        if "course_code" in col_map and "grade" in col_map:
            return col_map
        return None

    @classmethod
    def _parse_table_row(
        cls, row: list, col_map: Dict[str, int]
    ) -> Optional[CourseRecord]:
        """Parse a table row using identified column mapping."""
        try:
            course_code = row[col_map["course_code"]].strip()
            if not cls.COURSE_CODE_RE.match(course_code):
                return None

            grade = row[col_map.get("grade", -1)].strip()
            if not cls.GRADE_RE.match(grade):
                return None

            credits = 3.0
            if "credits" in col_map:
                try:
                    credits = float(row[col_map["credits"]] or 3)
                except (ValueError, TypeError):
                    credits = 3.0

            semester = "Unknown"
            if "semester" in col_map:
                semester = row[col_map["semester"]].strip() or "Unknown"

            return CourseRecord(
                course_code=course_code,
                credits=credits,
                grade=grade,
                semester=semester,
            )
        except (IndexError, ValueError, TypeError):
            return None

    @classmethod
    def _parse_row_positional(cls, row: list) -> Optional[CourseRecord]:
        """Parse a row without headers using positional guessing."""
        if len(row) < 3:
            return None

        course_code = None
        grade = None
        credits = 3.0
        semester = "Unknown"

        for cell in row:
            if not course_code and cls.COURSE_CODE_RE.match(cell):
                course_code = cell
            elif not grade and cls.GRADE_RE.match(cell):
                grade = cell
            elif cls.SEMESTER_RE.search(cell):
                sem_m = cls.SEMESTER_RE.search(cell)
                semester = f"{sem_m.group(1).title()} {sem_m.group(2)}"
            else:
                try:
                    val = float(cell)
                    if 0 <= val <= 6:
                        credits = val
                except ValueError:
                    pass

        if course_code and grade:
            try:
                return CourseRecord(
                    course_code=course_code,
                    credits=credits,
                    grade=grade,
                    semester=semester,
                )
            except ValueError:
                return None
        return None
